import streamlit as st
from docx import Document
import random
import io
import zipfile
import re

# --- CẤU HÌNH TRANG WEB ---
st.set_page_config(page_title="Trộn Đề Trắc Nghiệm", page_icon="📝")

# --- HÀM XỬ LÝ LOGIC ---
def parse_questions(doc_file):
    """Đọc file Word và tách câu hỏi, đáp án."""
    doc = Document(doc_file)
    questions = []
    current_q = None
    
    # Regex nhận diện câu hỏi và đáp án
    q_pattern = re.compile(r'^(Câu\s+\d+|Câu\s+hỏi\s+\d+|Bài\s+\d+)', re.IGNORECASE)
    opt_pattern = re.compile(r'^([#]?[a-dA-D])[\.\)]\s*(.*)')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if q_pattern.match(text) or (current_q is None and not opt_pattern.match(text)):
            if current_q:
                questions.append(current_q)
            current_q = {"content": text, "options": [], "correct": None}
        
        elif current_q:
            match = opt_pattern.match(text)
            if match:
                prefix = match.group(1).lower()
                content = match.group(2)
                is_correct = prefix.startswith('#')
                
                current_q["options"].append({
                    "text": content,
                    "is_correct": is_correct
                })
            else:
                current_q["content"] += "\n" + text

    if current_q:
        questions.append(current_q)
    return questions

def generate_mixed_exam(questions, exam_code):
    """Trộn câu hỏi và tạo file Word mới."""
    doc = Document()
    doc.add_heading(f'ĐỀ THI TRẮC NGHIỆM - MÃ ĐỀ {exam_code}', 0)
    
    mixed_qs = questions.copy()
    random.shuffle(mixed_qs) 
    
    answer_key = {} 
    labels = ['A', 'B', 'C', 'D']
    
    for idx, q in enumerate(mixed_qs, 1):
        doc.add_paragraph(f"Câu {idx}: {q['content'].split(':', 1)[-1].strip() if ':' in q['content'] else q['content']}")
        options = q['options'].copy()
        random.shuffle(options)
        
        for i, opt in enumerate(options):
            label = labels[i]
            doc.add_paragraph(f"{label}. {opt['text']}")
            if opt['is_correct']:
                answer_key[idx] = label
        doc.add_paragraph("") 
        
    return doc, answer_key

def create_answer_sheet(all_keys):
    """Tạo file đáp án tổng hợp."""
    doc = Document()
    doc.add_heading('BẢNG ĐÁP ÁN TỔNG HỢP', 0)
    
    for code, keys in all_keys.items():
        doc.add_heading(f'Mã đề: {code}', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = 'Câu'
        table.rows[0].cells[1].text = 'Đáp án'
        
        sorted_keys = dict(sorted(keys.items()))
        for q_num, ans in sorted_keys.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(q_num)
            row_cells[1].text = ans
        doc.add_paragraph("\n")
    return doc

# --- GIAO DIỆN WEB ---
st.title("📝 Công Cụ Trộn Đề Trắc Nghiệm Online")
st.markdown("---")

with st.expander("📖 Xem hướng dẫn cấu trúc file Word", expanded=True):
    st.markdown("""
    **Quy ước soạn thảo file Word (.docx):**
    1. **Câu hỏi:** Bắt đầu bằng chữ `Câu`.
    2. **Đáp án:** a., b., c., d.
    3. **Đáp án đúng:** Thêm dấu `#` trước ký tự (Ví dụ: `#a.`, `#c.`).
    """)

uploaded_file = st.file_uploader("Tải lên file Word đề gốc (.docx)", type=['docx'])

if uploaded_file is not None:
    st.success(f"Đã nhận file: {uploaded_file.name}")
    
    if st.button("🚀 Bắt đầu trộn đề"):
        with st.spinner('Đang xử lý...'):
            try:
                questions = parse_questions(uploaded_file)
                if not questions:
                    st.error("Lỗi: Không tìm thấy câu hỏi đúng định dạng!")
                else:
                    st.info(f"Đã tìm thấy {len(questions)} câu hỏi. Đang tạo 4 mã đề...")
                    
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                        all_keys = {}
                        for code in [101, 102, 103, 104]:
                            doc_exam, keys = generate_mixed_exam(questions, code)
                            all_keys[code] = keys
                            exam_io = io.BytesIO()
                            doc_exam.save(exam_io)
                            zip_file.writestr(f"De_Thi_{code}.docx", exam_io.getvalue())
                        
                        doc_answers = create_answer_sheet(all_keys)
                        ans_io = io.BytesIO()
                        doc_answers.save(ans_io)
                        zip_file.writestr("Dap_An_Tong_Hop.docx", ans_io.getvalue())

                    st.markdown("### ✅ Xử lý xong!")
                    st.download_button(
                        label="📥 Tải về bộ đề (.zip)",
                        data=zip_buffer.getvalue(),
                        file_name="Ket_Qua_Tron_De.zip",
                        mime="application/zip"
                    )
            except Exception as e:
                st.error(f"Có lỗi xảy ra: {e}")